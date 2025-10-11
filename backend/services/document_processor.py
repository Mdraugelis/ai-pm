"""
Document Processor Service
AI Product Manager Agent - Backend

Handles parsing and content extraction from various file formats.
Supports: PDF, DOCX, TXT, MD, JSON, and more.
"""

import json
import mimetypes
from pathlib import Path
from typing import Any, Dict, Optional, Tuple

import structlog

logger = structlog.get_logger(__name__)


# ============================================================================
# Document Processor
# ============================================================================


class DocumentProcessor:
    """
    Extract text content from various file formats

    Supports:
    - Text files: .txt, .md, .json, .yaml, .yml
    - PDF files: .pdf
    - Word documents: .docx, .doc
    - Rich text: .rtf

    Example:
        >>> processor = DocumentProcessor()
        >>> result = processor.extract_content("document.pdf")
        >>> print(result["content"])
    """

    def __init__(self):
        """Initialize document processor"""
        logger.info("DocumentProcessor initialized")

    def extract_content(
        self, file_path: str, file_bytes: Optional[bytes] = None
    ) -> Dict[str, Any]:
        """
        Extract content from file

        Args:
            file_path: Path to file (used for type detection)
            file_bytes: Optional file bytes if not reading from disk

        Returns:
            Dict with:
                - content: Extracted text content
                - file_type: Detected file type
                - mime_type: MIME type
                - metadata: Additional file metadata
                - status: "success" or "failed"
                - error: Error message if failed
        """
        logger.info("Extracting content from file", file_path=file_path)

        try:
            # Detect file type
            file_type, mime_type = self._detect_file_type(file_path)

            # Route to appropriate parser
            if file_type in ["txt", "md", "markdown"]:
                content = self._parse_text(file_path, file_bytes)

            elif file_type == "json":
                content = self._parse_json(file_path, file_bytes)

            elif file_type in ["yaml", "yml"]:
                content = self._parse_yaml(file_path, file_bytes)

            elif file_type == "pdf":
                content = self._parse_pdf(file_path, file_bytes)

            elif file_type in ["docx", "doc"]:
                content = self._parse_docx(file_path, file_bytes)

            elif file_type == "rtf":
                content = self._parse_rtf(file_path, file_bytes)

            else:
                # Try as plain text fallback
                logger.warning(
                    "Unknown file type, attempting text parsing",
                    file_type=file_type,
                )
                content = self._parse_text(file_path, file_bytes)

            # Build result
            result = {
                "status": "success",
                "content": content,
                "file_type": file_type,
                "mime_type": mime_type,
                "metadata": {
                    "filename": Path(file_path).name,
                    "file_extension": Path(file_path).suffix,
                    "content_length": len(content),
                },
            }

            logger.info(
                "Content extracted successfully",
                file_type=file_type,
                content_length=len(content),
            )

            return result

        except Exception as e:
            logger.error(
                "Content extraction failed",
                file_path=file_path,
                error=str(e),
                exc_info=True,
            )

            return {
                "status": "failed",
                "error": str(e),
                "file_type": "unknown",
                "content": "",
                "metadata": {"filename": Path(file_path).name},
            }

    # ========================================================================
    # File Type Detection
    # ========================================================================

    def _detect_file_type(self, file_path: str) -> Tuple[str, str]:
        """
        Detect file type from path

        Args:
            file_path: Path to file

        Returns:
            Tuple of (file_type, mime_type)
        """
        # Get extension
        extension = Path(file_path).suffix.lower().lstrip(".")

        # Detect MIME type
        mime_type, _ = mimetypes.guess_type(file_path)

        return extension, mime_type or "application/octet-stream"

    # ========================================================================
    # Format-Specific Parsers
    # ========================================================================

    def _parse_text(
        self, file_path: str, file_bytes: Optional[bytes] = None
    ) -> str:
        """
        Parse plain text file

        Args:
            file_path: Path to file
            file_bytes: Optional file bytes

        Returns:
            Text content
        """
        if file_bytes:
            return file_bytes.decode("utf-8", errors="replace")

        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    def _parse_json(
        self, file_path: str, file_bytes: Optional[bytes] = None
    ) -> str:
        """
        Parse JSON file and convert to readable text

        Args:
            file_path: Path to file
            file_bytes: Optional file bytes

        Returns:
            Formatted JSON as text
        """
        if file_bytes:
            data = json.loads(file_bytes.decode("utf-8"))
        else:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)

        # Convert to pretty-printed JSON
        return json.dumps(data, indent=2)

    def _parse_yaml(
        self, file_path: str, file_bytes: Optional[bytes] = None
    ) -> str:
        """
        Parse YAML file

        Args:
            file_path: Path to file
            file_bytes: Optional file bytes

        Returns:
            YAML content as text
        """
        try:
            import yaml

            if file_bytes:
                data = yaml.safe_load(file_bytes.decode("utf-8"))
            else:
                with open(file_path, "r", encoding="utf-8") as f:
                    data = yaml.safe_load(f)

            # Convert back to YAML for consistent formatting
            return yaml.dump(data, default_flow_style=False)

        except ImportError:
            logger.warning("PyYAML not installed, parsing as text")
            return self._parse_text(file_path, file_bytes)

    def _parse_pdf(
        self, file_path: str, file_bytes: Optional[bytes] = None
    ) -> str:
        """
        Parse PDF file

        Args:
            file_path: Path to file
            file_bytes: Optional file bytes

        Returns:
            Extracted text content
        """
        try:
            import pypdf

            if file_bytes:
                from io import BytesIO

                pdf_file = BytesIO(file_bytes)
            else:
                pdf_file = open(file_path, "rb")

            reader = pypdf.PdfReader(pdf_file)

            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"--- Page {page_num + 1} ---\n{text}")

            if not isinstance(file_bytes, bytes):
                pdf_file.close()

            return "\n\n".join(text_parts)

        except ImportError:
            logger.error("pypdf not installed, cannot parse PDF")
            raise ValueError(
                "PDF parsing requires pypdf library: pip install pypdf"
            )

    def _parse_docx(
        self, file_path: str, file_bytes: Optional[bytes] = None
    ) -> str:
        """
        Parse DOCX file

        Args:
            file_path: Path to file
            file_bytes: Optional file bytes

        Returns:
            Extracted text content
        """
        try:
            from docx import Document

            if file_bytes:
                from io import BytesIO

                doc = Document(BytesIO(file_bytes))
            else:
                doc = Document(file_path)

            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)

            # Combine
            all_text = paragraphs
            if table_text:
                all_text.append("\n--- Tables ---")
                all_text.extend(table_text)

            return "\n\n".join(all_text)

        except ImportError:
            logger.error("python-docx not installed, cannot parse DOCX")
            raise ValueError(
                "DOCX parsing requires python-docx library: pip install python-docx"
            )

    def _parse_rtf(
        self, file_path: str, file_bytes: Optional[bytes] = None
    ) -> str:
        """
        Parse RTF file

        Args:
            file_path: Path to file
            file_bytes: Optional file bytes

        Returns:
            Extracted text content
        """
        try:
            from striprtf.striprtf import rtf_to_text

            if file_bytes:
                rtf_content = file_bytes.decode("utf-8", errors="replace")
            else:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    rtf_content = f.read()

            return rtf_to_text(rtf_content)

        except ImportError:
            logger.warning("striprtf not installed, parsing as text")
            return self._parse_text(file_path, file_bytes)

    # ========================================================================
    # Validation and Cleanup
    # ========================================================================

    def validate_content(self, content: str, min_length: int = 10) -> bool:
        """
        Validate extracted content

        Args:
            content: Extracted content
            min_length: Minimum required length

        Returns:
            True if valid, False otherwise
        """
        if not content:
            return False

        if len(content.strip()) < min_length:
            logger.warning(
                "Content too short",
                length=len(content),
                min_length=min_length,
            )
            return False

        return True

    def clean_content(self, content: str) -> str:
        """
        Clean and normalize extracted content

        Args:
            content: Raw extracted content

        Returns:
            Cleaned content
        """
        # Remove excessive whitespace
        lines = content.split("\n")
        cleaned_lines = []

        for line in lines:
            line = line.strip()
            if line:
                cleaned_lines.append(line)

        # Join with single newlines
        cleaned = "\n".join(cleaned_lines)

        # Remove excessive blank lines (more than 2 consecutive)
        import re

        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)

        return cleaned


# ============================================================================
# Module Exports
# ============================================================================

__all__ = ["DocumentProcessor"]
